import io, os, re, tempfile, streamlit as st

# === DOCX -> EPUB 3 (from *converted* DOCX bytes) ===
def docx_bytes_to_epub3(docx_bytes: bytes, split_on_heading=True):
    """Convert DOCX bytes to a minimal EPUB 3 package (bytes)."""
    import zipfile, io, xml.etree.ElementTree as ET, html, datetime, uuid
    # Read DOCX XML
    zf = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    try:
        doc_xml = zf.read("word/document.xml")
    except KeyError:
        raise RuntimeError("DOCX missing word/document.xml")
    try:
        styles_xml = zf.read("word/styles.xml")
    except KeyError:
        styles_xml = None
    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = ET.fromstring(doc_xml)
    style_name_by_id = {}
    if styles_xml:
        sroot = ET.fromstring(styles_xml)
        for s in sroot.findall(".//w:style", NS):
            sid = s.get("{%s}styleId" % NS["w"])
            name_el = s.find(".//w:name", NS)
            if sid and name_el is not None and name_el.get("{%s}val" % NS["w"]):
                style_name_by_id[sid] = name_el.get("{%s}val" % NS["w"])
    chapters, current_title, current_paras = [], None, []
    def flush_chapter():
        if current_paras or current_title is not None:
            body_html = "".join(current_paras).strip() or "<p></p>"
            title = current_title or "Untitled"
            chapters.append((title, body_html))
    # paragraphs -> chapters
    for p in root.findall(".//w:p", NS):
        pStyle = None
        pPr = p.find("w:pPr", NS)
        if pPr is not None:
            ps = pPr.find("w:pStyle", NS)
            if ps is not None and ps.get("{%s}val" % NS["w"]):
                pStyle = ps.get("{%s}val" % NS["w"])
        style_name = style_name_by_id.get(pStyle, "") or ""
        texts = [t.text or "" for t in p.findall(".//w:t", NS)]
        text = "".join(texts).strip()
        if not text:
            current_paras.append("<p></p>")
            continue
        if split_on_heading and (style_name.lower().startswith("heading 1") or style_name.lower().startswith("heading 2")):
            flush_chapter()
            current_title, current_paras = html.escape(text), []
        else:
            current_paras.append("<p>%s</p>" % html.escape(text))
    flush_chapter()
    if not chapters:
        body_texts = []
        for p in root.findall(".//w:p", NS):
            texts = [t.text or "" for t in p.findall(".//w:t", NS)]
            t = "".join(texts).strip()
            if t:
                body_texts.append("<p>%s</p>" % html.escape(t))
        chapters = [("Document", "".join(body_texts) or "<p></p>")]
    book_id = "urn:uuid:" + str(uuid.uuid4())
    now = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    def make_opf(n_items):
        manifest_items, spine_items = [], []
        for i in range(n_items):
            manifest_items.append('<item id="c%s" href="chapter-%s.xhtml" media-type="application/xhtml+xml"/>' % (i+1, i+1))
            spine_items.append('<itemref idref="c%s"/>' % (i+1))
        manifest = "\n      ".join(manifest_items + ['<item id="nav" href="nav.xhtml" properties="nav" media-type="application/xhtml+xml"/>'])
        spine = "\n      ".join(spine_items)
        return '<?xml version="1.0" encoding="utf-8"?>\n<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="pub-id">\n  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">\n    <dc:identifier id="pub-id">%s</dc:identifier>\n    <dc:title>Converted Document</dc:title>\n    <dc:language>en</dc:language>\n    <meta property="dcterms:modified">%s</meta>\n  </metadata>\n  <manifest>\n      %s\n  </manifest>\n  <spine>\n      %s\n  </spine>\n</package>' % (book_id, now, manifest, spine)
    def make_nav(chapters):
        lis = []
        for i, (title, _) in enumerate(chapters, 1):
            t = title or "Chapter %s" % i
            lis.append('<li><a href="chapter-%s.xhtml">%s</a></li>' % (i, t))
        ol = "\n        ".join(lis)
        return '<?xml version="1.0" encoding="utf-8"?>\n<!DOCTYPE html>\n<html xmlns="http://www.w3.org/1999/xhtml">\n  <head><title>Table of Contents</title><meta charset="utf-8"/></head>\n  <body>\n    <nav epub:type="toc" id="toc">\n      <h2>Contents</h2>\n      <ol>\n        %s\n      </ol>\n    </nav>\n  </body>\n</html>' % ol
    def make_chapter(title, body):
        return '<?xml version="1.0" encoding="utf-8"?>\n<!DOCTYPE html>\n<html xmlns="http://www.w3.org/1999/xhtml">\n  <head><title>%s</title><meta charset="utf-8"/></head>\n  <body>\n    <h1>%s</h1>\n    %s\n  </body>\n</html>' % (title, title, body)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w") as ez:
        zinfo = zipfile.ZipInfo("mimetype")
        zinfo.compress_type = zipfile.ZIP_STORED
        ez.writestr(zinfo, "application/epub+zip")
        ez.writestr("META-INF/container.xml",
                    '<?xml version="1.0"?>\n<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\n  <rootfiles>\n    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>\n  </rootfiles>\n</container>')
        chapters_count = len(chapters)
        ez.writestr("OEBPS/content.opf", make_opf(chapters_count))
        ez.writestr("OEBPS/nav.xhtml", make_nav(chapters))
        for i, (title, body) in enumerate(chapters, 1):
            ez.writestr("OEBPS/chapter-%s.xhtml" % i, make_chapter(title, body))
    return out.getvalue()

# === ACBD drop-cap fixer (refined, PDF->DOCX only; no UI changes) ===
# Rules per user:
# • A = single large glyph (≥ 1.5× paragraph median size), usually letter + space.
# • B = subsequent normal-size runs after A up to C-start.
# • C = starts at the first ALL-CAPS word (≥2 letters), which may be split across runs
#   and may occur in the same or later paragraphs; C continues until a paragraph
#   that contains <w:widowControl/> is encountered (exclusive).
# • D = the widowControl paragraph (left untouched).
# • Reorder: current paragraph text becomes A + C + " " + B. Repeat until stable.
# Diagnostics: set ACBD_DIAG = True
ACBD_GLOBAL_MEDIAN_SIZE = None  # for per-paragraph prints.
ACBD_DIAG = False
ACBD_GLOBAL_MEDIAN_SIZE = None
ACBD_LOG = []

def _acbd_log(msg: str):
    ACBD_LOG.append(str(msg))
    try:
        print(str(msg))
    except Exception:
        pass

import statistics as _acbd_stats
import re as _acbd_re

#def _acbd_is_letter_space(txt: str) -> bool:
#    """Return True if txt is exactly: one uppercase A–Z followed by exactly one space (regular or NBSP)."""
#    if txt is None:
#        return False
#    # Preserve spaces; remove control chars
#    t = txt.replace("\u00A0", "\u0020")  # NBSP -> space
#    # Accept exactly two chars: [A-Z][space]
#    return bool(_acbd_re.fullmatch(r"[A-Z] ", t))


def _acbd_pt(val, default=None):
    try:
        return float(val.pt) if hasattr(val, "pt") else (float(val) if val is not None else default)
    except Exception:
        return default

def _acbd_run_size_pt(run, para, default=11.0):
    # Prefer python-docx size if present; else pull from XML (<w:sz> or <w:szCs>), else paragraph style
    sz = _acbd_pt(getattr(run.font, "size", None))
    if sz is not None:
        return sz
    try:
        el = run._element
        vals = el.xpath(".//w:rPr/w:sz/@w:val | .//w:rPr/w:szCs/@w:val", namespaces=el.nsmap)
        if vals:
            try:
                return float(vals[0]) / 2.0  # half-points -> points
            except Exception:
                pass
        # Also check paragraph-level rPr if present
        pel = para._element
        pvals = pel.xpath(".//w:pPr/w:rPr/w:sz/@w:val | .//w:pPr/w:rPr/w:szCs/@w:val", namespaces=pel.nsmap)
        if pvals:
            try:
                return float(pvals[0]) / 2.0
            except Exception:
                pass
        psz = _acbd_pt(para.style.font.size, None)
        if psz is not None:
            return psz
    except Exception:
        pass
    return default

def _acbd_run_text(run):
    # Concatenate all <w:t> in this run
    try:
        return "".join(t.text or "" for t in run._element.xpath(".//w:t", namespaces=run._element.nsmap))
    except Exception:
        return getattr(run, "text", "") or ""

def _acbd_first_caps_token_across_runs(doc, start_para, start_run):
    """
    Scan (para, run) sequence starting at (start_para, start_run) to find earliest ALL-CAPS token (>=2 letters).
    Allows tokens split across adjacent runs.
    Returns (para_index, run_index, char_offset_in_run) for the token start, or None.
    """
    paras = doc.paragraphs
    token = ""
    started = False
    for pi in range(start_para, len(paras)):
        runs = paras[pi].runs
        ri0 = start_run if pi == start_para else 0
        for ri in range(ri0, len(runs)):
            txt = _acbd_run_text(runs[ri])
            for ci, ch in enumerate(txt):
                if ch.isalpha():
                    if not started:
                        token = ch
                        start_loc = (pi, ri, ci)
                        started = True
                    else:
                        token += ch
                else:
                    if started and len(token) >= 2 and token.upper() == token:
                        return start_loc
                    token = ""
                    started = False
            # End of run boundary acts like a separator; evaluate token so far
            if started and len(token) >= 2 and token.upper() == token:
                return start_loc
    return None

def _acbd_para_has_widowcontrol(para):
    try:
        el = para._element
        return bool(el.xpath(".//w:widowControl", namespaces=el.nsmap))
    except Exception:
        return False

def _acbd_find_widowcontrol_forward(doc, start_para):
    """Return index of first paragraph >= start_para that contains <w:widowControl/>, else None."""
    for pi in range(start_para, len(doc.paragraphs)):
        if _acbd_para_has_widowcontrol(doc.paragraphs[pi]):
            return pi
    return None



def _acbd_doc_global_median_size(doc, default=12.0):
    """Compute a global median font size (points) over ALL runs in the document using w:sz/w:szCs/xml fallbacks."""
    sizes = []
    for p in doc.paragraphs:
        for r in p.runs:
            try:
                # Reuse existing sizing function if present
                sz = _acbd_run_size_pt(r, p, default=None)
            except Exception:
                sz = None
            if sz is not None:
                sizes.append(sz)
    if not sizes:
        return default
    try:
        import statistics as _stats
        return float(_stats.median(sizes))
    except Exception:
        return sum(sizes)/len(sizes)

#def _acbd_para_median_size(para):
#    sizes = [_acbd_run_size_pt(r, para) for r in para.runs]
#    sizes = [s for s in sizes if s is not None]
#    if not sizes:
#        return 11.0
#    try:
#        return float(_acbd_stats.median(sizes))
#    except Exception:
#        return sum(sizes)/len(sizes)

def _acbd_fix_once_in_paragraph(doc, p_index):
    paras = doc.paragraphs
    if p_index < 0 or p_index >= len(paras):
        return False
    p = paras[p_index]
    runs = list(p.runs)
    if not runs:
        return False

    # Gather run info (size, text) for this paragraph
    run_info = []
    for i, r in enumerate(runs):
        txt = _acbd_run_text(r)
        sz = _acbd_run_size_pt(r, p)
        run_info.append((i, sz, txt))

    # Use GLOBAL median as requested
    majority = (ACBD_GLOBAL_MEDIAN_SIZE if ACBD_GLOBAL_MEDIAN_SIZE is not None else 12.0)
    threshold = 1.5 * majority
    max_size = max((s for _, s, _ in run_info if s is not None), default=majority)

    if ACBD_DIAG:
        _acbd_log(f"[ACBD] p={p_index}: sizes(med={majority:.1f}, thr={threshold:.1f}, max={max_size:.1f})")
        # Show top runs by size
        top_runs = sorted(run_info, key=lambda t: (t[1] or -1), reverse=True)[:5]
        for (ri, rsz, rtxt) in top_runs:
            preview = (rtxt or '')[:30]
            _acbd_log(f"    [run {ri}] sz={rsz} text={repr(preview)}")

    # Primary A detection: single uppercase letter + a space (or NBSP) in this run,
    # size >= threshold; OR the next run is pure whitespace
    A_idx = None
    A_char = None
    for i, sz, txt in run_info:
        if not txt:
            continue
        letters = [ch for ch in txt if ch.isalpha()]
        if len(letters) == 1 and letters[0].isupper() and sz is not None and sz >= threshold:
            ends_space = txt.endswith(" ") or txt.endswith("\u00A0")
            next_is_space = False
            if i + 1 < len(run_info):
                nxt_txt = run_info[i+1][2] or ""
                nxt_norm = nxt_txt.replace("\u00A0", " ")
                next_is_space = (nxt_norm.strip() == "")
            if ends_space or next_is_space:
                A_idx = i
                A_char = letters[0]
                if ACBD_DIAG:
                    _acbd_log(f"[ACBD] p={p_index}: A at run {i} (sz={sz}) ends_space={ends_space} next_space={next_is_space}")
                break

    if A_idx is None:
        # No suitable A in this paragraph
        _acbd_log(f"[ACBD] p={p_index}: no A (thr={threshold:.1f}, med={majority:.1f}, max={max_size:.1f})")
        return False

    # Find C-start across runs/paragraphs; stop only if widowControl encountered before any ALL-CAPS
    c_start_loc = _acbd_first_caps_token_across_runs(doc, p_index, A_idx+1)
    wc_idx = _acbd_find_widowcontrol_forward(doc, p_index+1)

    if wc_idx is not None and (c_start_loc is None or c_start_loc[0] >= wc_idx):
        _acbd_log(f"[ACBD] p={p_index}: widowControl@{wc_idx} before C-start; skip")
        return False
    if c_start_loc is None:
        _acbd_log(f"[ACBD] p={p_index}: no C-start found in document tail; skip")
        return False

    c_pi, c_ri, c_ci = c_start_loc

    # If no widowControl is found later, fall back to using the C-start paragraph as terminator
    if wc_idx is None:
        _acbd_log(f"[ACBD] p={p_index}: no widowControl found; FALLBACK to C within C-start paragraph")
        wc_idx = c_pi

    # Build B (text between A and C-start within this paragraph if C-start here, else all remaining runs)
    if c_pi == p_index:
        B_text = "".join(_acbd_run_text(runs[t]) for t in range(A_idx+1, c_ri)).strip()
    else:
        B_text = "".join(_acbd_run_text(runs[t]) for t in range(A_idx+1, len(runs))).strip()

    # Build C text from c_start to the paragraph just before widowControl
    C_parts = []
    c_runs = paras[c_pi].runs
    start_txt = _acbd_run_text(c_runs[c_ri])
    # Include from c_ci (char offset) onward in the starting run
    C_parts.append(start_txt[c_ci:] if c_ci < len(start_txt) else "")
    for t in range(c_ri+1, len(c_runs)):
        C_parts.append(_acbd_run_text(c_runs[t]))
    for pi in range(c_pi+1, wc_idx):
        C_parts.extend(_acbd_run_text(r) for r in paras[pi].runs)
    C_text = "".join(C_parts).strip()

    if not B_text or not C_text:
        _acbd_log(f"[ACBD] p={p_index}: empty B or C (B={len(B_text)}, C={len(C_text)}); skip")
        return False

    # Recompose current paragraph: A + C + " " + B
    new_text = (A_char.upper() + C_text).strip()
    if B_text:
        new_text += " " + B_text

    if new_text != (p.text or "").strip():
        _acbd_log(f"[ACBD] p={p_index}: REORDERED | A='{A_char}' | B[:30]='{B_text[:30]}' | C[:30]='{C_text[:30]}' | wc@{wc_idx} c@({c_pi},{c_ri},{c_ci})")
        p.text = new_text
        return True
    else:
        _acbd_log(f"[ACBD] p={p_index}: no change after recomposition")
        return False

def fix_dropcaps_acbd(doc, max_passes=80):
    global ACBD_GLOBAL_MEDIAN_SIZE
    ACBD_GLOBAL_MEDIAN_SIZE = _acbd_doc_global_median_size(doc)
    passes = 0
    while passes < max_passes:
        changes = 0
        for i in range(len(doc.paragraphs)):
            inner = 0
            while inner < 6 and _acbd_fix_once_in_paragraph(doc, i):
                changes += 1
                inner += 1
        _acbd_log(f"[ACBD] pass={passes} changes={changes} global_med={ACBD_GLOBAL_MEDIAN_SIZE}")
        if changes == 0:
            break
        passes += 1
    return doc


def acbd_write_log(sidecar_path=None):
    """
    Write ACBD diagnostics to a sidecar text file.
    If sidecar_path is None, default to '/mnt/data/ACBD_diagnostics.txt'.
    """
    path = sidecar_path or "/mnt/data/ACBD_diagnostics.txt"
    try:
        with open(path, "w", encoding="utf-8") as f:
            for line in ACBD_LOG:
                f.write(line.rstrip("\n") + "\n")
    except Exception as e:
        try:
            print(f"[ACBD] failed to write log: {e}")
        except Exception:
            pass
# === end refined ACBD fixer ===



try:
    from docx import Document
except Exception:
    Document = None

try:
    from pdf2docx import Converter as PDF2DOCXConverter
except Exception:
    PDF2DOCXConverter = None

_ASCII_CTRL = re.compile(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]')

def _drop_nonchars(s: str) -> str:
    out = []
    for ch in s:
        code = ord(ch)
        if 0xFDD0 <= code <= 0xFDEF or (code & 0xFFFE) == 0xFFFE:
            continue
        if 0xD800 <= code <= 0xDFFF:
            out.append('\uFFFD'); continue
        out.append(ch)
    return ''.join(out)

def _xml10_filter(text: str) -> str:
    if not text:
        return text
    out = []
    for ch in text:
        code = ord(ch)
        if code in (0x9, 0xA, 0xD) or 0x20 <= code <= 0xD7FF or 0xE000 <= code <= 0xFFFD or 0x10000 <= code <= 0x10FFFF:
            out.append(ch)
    return ''.join(out)

def sanitize_for_docx(text: str) -> str:
    if not text:
        return text
    text = _ASCII_CTRL.sub('', text)
    text = _drop_nonchars(text)
    return _xml10_filter(text)

def _detect_primary_style(text: str) -> str:
    if not text:
        return "UNKNOWN"
    singles_open = len(re.findall(r'(^|[\s(\[{<])‘', text))
    doubles_open = len(re.findall(r'(^|[\s(\[{<])“', text))
    singles_total = text.count("‘") + text.count("’")
    doubles_total = text.count("“") + text.count("”")
    if singles_open >= doubles_open * 1.5 and singles_open >= 4:
        return "UK"
    if doubles_open >= singles_open * 1.2 and doubles_open >= 4:
        return "US"
    if doubles_total > singles_total * 1.2 and doubles_open >= 2:
        return "US"
    if singles_total > doubles_total * 1.5 and singles_open >= 2:
        return "UK"
    return "UNKNOWN"

def normalize_quotes_to_us(text: str) -> str:
    if not text:
        return text
    APOS = "<<APOS>>"
    text = re.sub(r"(?<=\w)[’'](?=\w)", APOS, text)
    style = _detect_primary_style(text)
    if style == "UK":
        OPEN_S, CLOSE_S, OPEN_D, CLOSE_D = "<<OPEN_S>>", "<<CLOSE_S>>", "<<OPEN_D>>", "<<CLOSE_D>>"
        t = (text.replace("‘", OPEN_S)
                 .replace("’", CLOSE_S)
                 .replace("“", OPEN_D)
                 .replace("”", CLOSE_D))
        t = re.sub(r'(?<=\w)'+re.escape(CLOSE_S)+r'(?=\w)', APOS, t)
        for w in ("em","cause","til","tis","twas","sup","round","clock"):
            t = re.sub(r'\b'+re.escape(CLOSE_S)+w+r'\b', APOS+w, t, flags=re.IGNORECASE)
        t = re.sub(re.escape(CLOSE_S)+r'(?=\d{2}s\b)', APOS, t)
        t = (t.replace(OPEN_S,"“").replace(CLOSE_S,"”").replace(OPEN_D,"‘").replace(CLOSE_D,"’"))
        text = t
    else:
        def smarten_line(line: str) -> str:
            out, open_d = [], True
            for ch in line:
                if ch == '"':
                    out.append("“" if open_d else "”"); open_d = not open_d
                elif ch == "'":
                    out.append("’")
                else:
                    out.append(ch)
            return "".join(out)
        text = "\n".join(smarten_line(ln) for ln in text.split("\n"))
    return text.replace(APOS, "’")

def convert_docx_runs_to_us(doc: Document) -> None:
    for p in doc.paragraphs:
        for r in p.runs:
            r.text = normalize_quotes_to_us(sanitize_for_docx(r.text))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = normalize_quotes_to_us(sanitize_for_docx(r.text))

def _remove_global_shapes_all_parts(doc: Document) -> None:
    """
    Delete drawing/pict/object/sym/txbx/wsp elements from the main document and all related parts
    (headers/footers), then remove empty runs and paragraphs.
    """
    pkg = doc.part.package
    for part in pkg.parts:
        elt = getattr(part, 'element', None)
        if elt is None:
            continue
        # 1) Remove drawings/picts/objects/symbols and deep textbox containers
        nodes = list(elt.xpath(
            './/*[local-name()="drawing" or local-name()="pict" or local-name()="object" or local-name()="sym" or local-name()="wsp" or local-name()="txbx" or local-name()="txbxContent"]'
        ))
        for n in nodes:
            parent = n.getparent()
            if parent is not None:
                parent.remove(n)
        # 2) Remove empty runs
        for r in list(elt.xpath('.//*[local-name()="r"]')):
            has_text = bool(r.xpath('.//*[local-name()="t" and normalize-space(text())]'))
            has_children = len(r) > 0
            if not has_text and not has_children:
                parent = r.getparent()
                if parent is not None:
                    parent.remove(r)
        # 3) Remove paragraphs that are now empty or whitespace-only
        for p in list(elt.xpath('.//*[local-name()="p"]')):
            has_text = bool(p.xpath('.//*[local-name()="t" and normalize-space(text())]'))
            has_draw = bool(p.xpath('.//*[local-name()="drawing" or local-name()="pict" or local-name()="object" or local-name()="sym" or local-name()="wsp" or local-name()="txbx" or local-name()="txbxContent"]'))
            if not has_text and not has_draw:
                parent = p.getparent()
                if parent is not None:
                    parent.remove(p)

def convert_docx_bytes_to_us(docx_bytes: bytes) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx required.")
    doc = Document(io.BytesIO(docx_bytes))
    convert_docx_runs_to_us(doc)
    out = io.BytesIO(); doc.save(out)
    return out.getvalue()

def pdf_bytes_to_docx_using_pdf2docx(pdf_bytes: bytes) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx required.")
    if PDF2DOCXConverter is None:
        raise RuntimeError("pdf2docx required.")
    with tempfile.TemporaryDirectory() as tmpd:
        pdf_path = os.path.join(tmpd, "in.pdf")
        out_path = os.path.join(tmpd, "out.docx")
        with open(pdf_path, "wb") as f: f.write(pdf_bytes)
        cv = PDF2DOCXConverter(pdf_path)
        cv.convert(out_path, start=0, end=None)
        cv.close()
        doc = Document(out_path)
        fix_dropcaps_acbd(doc)
        acbd_write_log()

        # 1) Deep removal across all parts (fix persistent squares)
        _remove_global_shapes_all_parts(doc)

        # 2) Run-level cleanup and cautious mid-sentence blank removal
        paras = doc.paragraphs
        for i, p in enumerate(paras):
            for r in p.runs:
                if r.text:
                    r.text = (r.text.replace("\uFFFC","")
                                   .replace("\u00A0"," ")
                                   .replace("\u000c",""))
            if p.text.strip() in {"", "\u00A0"} and 0 < i < len(paras)-1:
                prev = paras[i-1].text.strip()
                nxt  = paras[i+1].text.strip()
                if prev and nxt and not re.search(r'[.!?]"?$', prev):
                    p.text = ""

        # 3) Normalize quotes to US
        convert_docx_runs_to_us(doc)

        buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

st.set_page_config(page_title="Quote Style Converter (Global Clean v3)", page_icon="📝", layout="centered")

CSS = """:root {
  --primary-color: #008080;      /* Teal */
  --primary-hover: #007070;
  --background-color: #fdfdfd;
  --text-color: #222222;
  --card-background: #ffffff;
  --card-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
  --border-radius: 10px;
  --font-family: 'Avenir', sans-serif;
  --accent-color: #ff9900;
}

/* Global Styles */
body {
  background-color: var(--background-color);
  font-family: var(--font-family);
  color: var(--text-color);
  margin: 0;
  padding: 0;
}

h1, h2, h3, h4, h5, h6 {
  color: var(--text-color);
  font-weight: 700;
  margin-bottom: 0.5em;
}

/* Button Styles */
div.stButton > button {
  background-color: var(--primary-color);
  color: #ffffff;
  border: none;
  padding: 0.75em 1.25em;
  border-radius: var(--border-radius);
  cursor: pointer;
  transition: background-color 0.3s ease, transform 0.2s;
}

div.stButton > button:hover {
  background-color: var(--primary-hover);
  transform: translateY(-2px);
}

/* Card/Container Styling */
.custom-container {
  background: var(--card-background);
  padding: 2em;
  border-radius: var(--border-radius);
  box-shadow: var(--card-shadow);
  margin-bottom: 2em;
}

.css-1d391kg {
  background: var(--card-background);
  padding: 1em;
  border-radius: var(--border-radius);
  box-shadow: var(--card-shadow);
}

/* Form Element Styling */
input, select, textarea {
  border: 1px solid #ddd;
  border-radius: 4px;
  padding: 0.5em;
  font-size: 1em;
}

input:focus, select:focus, textarea:focus {
  outline: none;
  border-color: var(--primary-color);
  box-shadow: 0 0 5px rgba(0, 128, 128, 0.3);
}

/* Enforce font in uploader */
.stFileUploader, .stFileUploader label, .stFileUploader div, .stFileUploader button, .stFileUploader *,
[data-testid="stFileUploader"], [data-testid="stFileUploader"] *, [data-testid="stFileUploadDropzone"], [data-testid="stFileUploadDropzone"] *,
input[type="file"] {
  font-family: var(--font-family) !important;
}

"""
st.markdown("<style>\n"+CSS+"\n</style>", unsafe_allow_html=True)

st.title("UK to US Quote Converter with Optional PDF to DOCX Conversion")
st.write("Please upload a docx using single-quotes dialogue for conversion to double-quotes dialogue, or upload a PDF of either type for conversion to double-quotes dialogue in a docx.")

uploaded = st.file_uploader(
    "Upload DOCX (single-quotes) or PDF",
    type=["docx", "pdf"],
    accept_multiple_files=False,
    key="file",
    label_visibility="collapsed"
)



if uploaded is not None:
    # Show a simple file summary and a Convert button
    st.write(f"Selected file: **{uploaded.name}**")
    if st.button("Convert"):
        name_lower = uploaded.name.lower()
        if name_lower.endswith(".docx"):

            if Document is None:

                st.error("python-docx not available; cannot process DOCX.")

            else:

                try:

                    raw = uploaded.read()

                    try:

                        out_bytes = docx_bytes_to_us_quotes(raw)

                    except NameError:

                        out_bytes = convert_docx_bytes_to_us(raw)

                    st.success("Converted. Download below.")

                    st.download_button("Download File", out_bytes,

                        file_name=uploaded.name,

                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                    try:

                        epub_bytes = docx_bytes_to_epub3(out_bytes)

                        st.download_button("Download EPUB 3", epub_bytes,

                            file_name=uploaded.name.rsplit(".",1)[0] + ".epub",

                            mime="application/epub+zip")

                    except Exception as _epub_err:

                        st.warning(f"EPUB generation skipped: {_epub_err}")

                except Exception as e:

                    st.error(f"Conversion failed: {e}")

        elif name_lower.endswith(".pdf"):
            if PDF2DOCXConverter is None:
                st.error("pdf2docx not available; cannot convert PDF to DOCX.")
            else:
                try:
                    out_bytes = pdf_bytes_to_docx_using_pdf2docx(uploaded.read())
                    st.success("Converted. Download below.")
                    base = uploaded.name.rsplit(".",1)[0]
                    st.download_button("Download File", out_bytes,
                        file_name=base + ".docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"Conversion failed: {e}")
        else:
            st.error("Unsupported file type. Please upload a .docx or .pdf.")